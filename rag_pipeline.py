# -- coding: utf-8 --

import torch
import os
import json
import warnings
import re
import shutil
import time
import uuid
from tqdm import tqdm
import requests
import PyPDF2
import pandas as pd
from bs4 import BeautifulSoup
from typing import List, Dict, Any, Tuple

# LangChain & ML Imports
from langchain.prompts import PromptTemplate
from langchain.schema.runnable import RunnablePassthrough, RunnableLambda, RunnableParallel
from langchain.schema.output_parser import StrOutputParser
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# --- NEW IMPORTS for Hybrid Retrieval ---
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.cross_encoders import HuggingFaceCrossEncoder # For the reranker model

# For .doc conversion
import win32com.client

# For DOCX
try:
    import docx
except ImportError:
    print("docx package not found. Please install it using 'pip install python-docx'")

warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning, module='torch.utils. deel_ Državna')
warnings.filterwarnings("ignore", message="Can't initialize NVML")

print("Imports successful.")

# Configuration
print("--- Configuring Paths ---")
DATA_ROOT = os.path.abspath("./Askalma")
WORKING_DIR = os.path.abspath("temp")
os.makedirs(WORKING_DIR, exist_ok=True) # General temp for things like doc conversion

BOT_NAME = "AskAlma"
CREATOR_INFO = "I was built by Amartya Singh, Abhishek Bansal, and Aditya Bagri."
PURPOSE = f"I am {BOT_NAME}, an AI assistant for the IIITD college website documents."

# --- Chunking Config for BM25 Corpus ---
CORPUS_CHUNK_SIZE = 1500
CORPUS_CHUNK_OVERLAP = 300

# --- Hybrid Retriever Config ---
BM25_K_CANDIDATES = 30  # How many documents BM25 retrieves initially
CROSS_ENCODER_MODEL_NAME = "cross-encoder/ms-marco-MiniLM-L-6-v2" # Smaller, faster
# CROSS_ENCODER_MODEL_NAME = "cross-encoder/ms-marco-MiniLM-L-12-v2" # Larger, potentially better
RERANK_TOP_N = 10      # How many documents to keep after reranking

HISTORY_TURNS_FOR_RETRIEVAL = 4 # How many turns of chat history to include in the contextual query

DEVICE = "cuda" if torch.cuda.is_available() else "cpu" # For potential PyTorch ops, not directly for CrossEncoder init
TORCH_DTYPE = torch.float16 if DEVICE == "cuda" else torch.float32
print(f"Using device for potential PyTorch ops (e.g., embeddings if used): {DEVICE}")
print(f"Using dtype for potential PyTorch ops: {TORCH_DTYPE}")
print("Configuration loaded.")

# --- Data Loading Functions ---
def clean_text(text: str) -> str:
    if not isinstance(text, str): text = str(text)
    text = re.sub(r'\s+', ' ', text).lower()
    text = re.sub(r'[^a-zA-Z0-9\s.,!?-]', '', text)
    text = re.sub(r'nan', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\.{2,}', '.', text)
    return text.strip()

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                try: reader.decrypt('')
                except: print(f"Warning: Could not decrypt PDF {os.path.basename(file_path)}"); return ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text: text += page_text + "\n"
    except Exception as e: print(f"    Error reading PDF {os.path.basename(file_path)}: {e}")
    return clean_text(text)

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc_obj = docx.Document(file_path)
        full_text = [para.text for para in doc_obj.paragraphs]
        text = '\n'.join(full_text)
    except Exception as e: print(f"    Error reading DOCX {os.path.basename(file_path)}: {e}")
    return clean_text(text)

def doc_to_text(doc_filepath: str, output_format="txt") -> str:
    text_content = ""
    word_instance = None 
    com_initialized = False
    try:
        win32com.client.pythoncom.CoInitialize()
        com_initialized = True
        word_instance = win32com.client.Dispatch("Word.Application")
        word_instance.Visible = False
        abs_doc_filepath = os.path.abspath(doc_filepath)
        temp_dir = os.path.join(WORKING_DIR, "doc_conversion_temp")
        os.makedirs(temp_dir, exist_ok=True)
        temp_output_filename = os.path.splitext(os.path.basename(doc_filepath))[0] + f"_{uuid.uuid4().hex[:6]}.txt"
        temp_output_filepath = os.path.join(temp_dir, temp_output_filename)
        abs_temp_output_filepath = os.path.abspath(temp_output_filepath)
        doc = None
        try:
            doc = word_instance.Documents.Open(abs_doc_filepath)
            if output_format.lower() == "txt": doc.SaveAs(abs_temp_output_filepath, FileFormat=2)
            else: print(f"Unsupported output format for .doc: {output_format}"); return ""
        finally:
            if doc: doc.Close(False)
        if word_instance: word_instance.Quit(); word_instance = None 
        time.sleep(0.1)
        with open(abs_temp_output_filepath, 'r', encoding='utf-8', errors='replace') as f: text_content = f.read()
        try: os.remove(abs_temp_output_filepath)
        except OSError as e: print(f"Warning: Could not remove temp file {abs_temp_output_filepath}: {e}")
        return clean_text(text_content)
    except Exception as e: print(f"    Error converting .doc '{os.path.basename(doc_filepath)}' to text: {e}"); return ""
    finally:
        if word_instance:
            try: word_instance.Quit()
            except: pass
        if com_initialized: win32com.client.pythoncom.CoUninitialize()

def excel_to_text(xlsx_filepath: str) -> str:
    try:
        xls = pd.ExcelFile(xlsx_filepath)
        text_parts = []
        if not xls.sheet_names: print(f"    Warning: No sheets in {os.path.basename(xlsx_filepath)}."); return ""
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            sheet_text_content = [" | ".join(str(cell) for cell in row if pd.notna(cell) and str(cell).strip()) 
                                  for _, row in df.iterrows() if any(pd.notna(cell) and str(cell).strip() for cell in row)]
            if sheet_text_content: text_parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text_content))
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text) if full_text.strip() else ""
    except Exception as e: print(f"    Error converting Excel '{os.path.basename(xlsx_filepath)}': {e}"); return ""

def extract_text_from_html(file_path: str) -> str:
    text = "";
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            soup = BeautifulSoup(file, 'html.parser'); text = soup.get_text(separator='\n', strip=True)
    except Exception as e: print(f"    Error reading HTML {os.path.basename(file_path)}: {e}")
    return clean_text(text)
def extract_text_from_table(file_path: str) -> str: return extract_text_from_html(file_path)
def extract_text_from_json(file_path: str) -> str:
    text = "";
    try:
        with open(file_path, 'r', encoding='utf-8') as f: data = json.load(f)
        if isinstance(data, list):
            text_parts = [(" ".join(f"{k}: {v}" for k, v in item.items() if v is not None) if isinstance(item, dict) else str(item)) for item in data]
            text = "\n".join(text_parts)
        elif isinstance(data, dict): text = " ".join(f"{k}: {v}" for k, v in data.items() if v is not None)
        else: text = json.dumps(data, indent=2)
    except Exception as e: print(f"    Error reading JSON {os.path.basename(file_path)}: {e}")
    return clean_text(text)
print("Data loading functions defined.")

def load_documents(data_root: str) -> List[Document]:
    print(f"Loading documents from data path: {data_root}...")
    all_docs = []; min_doc_length = 10
    data_sources = {
        "attachments": {"path": os.path.join(data_root, "attachments"), "handlers": {'.pdf': extract_text_from_pdf, '.docx': extract_text_from_docx, '.doc': doc_to_text, '.xlsx': excel_to_text, '.xls': excel_to_text}},
        "html": {"path": os.path.join(data_root, "html"), "handlers": {'.html': extract_text_from_html}},
        "tables": {"path": os.path.join(data_root, "tables"), "handlers": {'.html': extract_text_from_table}},
        "text_pdfs": {"path": os.path.join(data_root, "text_pdfs"), "handlers": {'.txt': lambda fp: clean_text(open(fp, 'r', encoding='utf-8', errors='replace').read())}},
        "course_data": {"path": os.path.join(data_root, "course_data"), "handlers": {'.json': extract_text_from_json}}}
    for source_name, config_val in data_sources.items(): # Renamed config to config_val
        dir_path = config_val["path"]
        if not os.path.isdir(dir_path): print(f"Warning: Directory not found - {dir_path}, skipping."); continue
        print(f"Processing directory: {source_name}...")
        for filename in tqdm(os.listdir(dir_path), desc=f"Loading {source_name}"):
            file_path = os.path.join(dir_path, filename)
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext in config_val["handlers"]:
                    text = config_val["handlers"][file_ext](file_path)
                    if text and len(text.split()) >= min_doc_length:
                        all_docs.append(Document(page_content=text, metadata={"source": source_name, "file": filename, "full_path": file_path}))
    factual_file = os.path.join(data_root, 'factual_data_spanbert.json')
    if os.path.exists(factual_file):
        print(f"Processing factual data file: {factual_file}...")
        try:
            with open(factual_file, 'r', encoding='utf-8') as f: factual_data = json.load(f)
            for item in tqdm(factual_data, desc="Loading facts"):
                if "facts" in item and isinstance(item["facts"], str):
                    for fact_line in item["facts"].split("\n"):
                        cleaned_fact = clean_text(fact_line)
                        if cleaned_fact and len(cleaned_fact.split()) >= min_doc_length:
                            all_docs.append(Document(page_content=cleaned_fact, metadata={"source": "facts_json", "file": os.path.basename(factual_file), "original_item_title": item.get("title", "N/A")}))
        except Exception as e: print(f"Error loading or processing factual data {factual_file}: {e}")
    else: print(f"Warning: Factual data file not found at {factual_file}, skipping.")
    print(f"\nLoaded a total of {len(all_docs)} documents.")
    from collections import Counter
    print("Document counts by source:", Counter(doc.metadata.get('source', 'unknown') for doc in all_docs))
    return all_docs

# --- Hybrid Retriever Initialization ---
def initialize_hybrid_retriever_components(
    all_loaded_docs: List[Document],
    corpus_chunk_size: int,
    corpus_chunk_overlap: int,
    bm25_k_candidates: int,
    cross_encoder_model_name: str,
    rerank_top_n: int
    # device_type: str # Removed as HuggingFaceCrossEncoder handles device automatically
) -> Tuple[BM25Retriever, CrossEncoderReranker]:
    if not all_loaded_docs:
        raise ValueError("No documents loaded for hybrid retriever.")
    print(f"Initializing Hybrid Retriever (BM25 -> CrossEncoder) components...")

    # 1. Chunk documents for the corpus
    print(f"Chunking {len(all_loaded_docs)} documents for BM25/Reranker corpus...")
    corpus_text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=corpus_chunk_size,
        chunk_overlap=corpus_chunk_overlap
    )
    corpus_chunks = corpus_text_splitter.split_documents(all_loaded_docs)
    print(f"Created {len(corpus_chunks)} corpus chunks.")
    if not corpus_chunks:
        raise ValueError("No corpus chunks created. Cannot initialize retrievers.")

    # 2. Initialize BM25Retriever
    print(f"Initializing BM25Retriever with k={bm25_k_candidates}...")
    try:
        bm25_retriever = BM25Retriever.from_documents(
            documents=corpus_chunks,
            k=bm25_k_candidates
        )
        print("BM25Retriever initialized.")
    except Exception as e:
        print(f"Error initializing BM25Retriever: {e}. Ensure 'rank_bm25' is installed ('pip install rank_bm25').")
        raise

    # 3. Initialize CrossEncoderReranker
    print(f"Initializing CrossEncoderReranker with model '{cross_encoder_model_name}' and top_n={rerank_top_n}...")
    try:
        cross_encoder_model = HuggingFaceCrossEncoder(
            model_name=cross_encoder_model_name
            # device parameter removed as it's not accepted here.
            # sentence-transformers will attempt to use GPU if available.
        )
        reranker = CrossEncoderReranker(model=cross_encoder_model, top_n=rerank_top_n)
        print("CrossEncoderReranker initialized.")
    except Exception as e:
        print(f"Error initializing CrossEncoderReranker: {e}. Ensure 'sentence-transformers' is installed.")
        raise
        
    return bm25_retriever, reranker

# --- LLM, Prompt, Chain functions ---
def load_persona_data(persona_file_path: str) -> tuple[str, dict, dict]:
    try:
        with open(persona_file_path, 'r', encoding='utf-8') as f: data = json.load(f)
        return data.get("persona", ""), data.get("faqs", {}), data.get("paths", {})
    except Exception as e: print(f"Error loading persona data from {persona_file_path}: {e}"); return "I am a helpful AI.", {}, {}
PERSONA_FILE_PATH = os.path.join(DATA_ROOT, "persona.json")

def llm_pipeline_function(chain_input_dict: Dict[str, Any]) -> str:
    raw_full_prompt = chain_input_dict.get("full_prompt_str", "")
    chat_history: List[Dict[str, str]] = chain_input_dict.get("chat_history", [])
    full_prompt_as_string = ""
    if hasattr(raw_full_prompt, 'to_string'): full_prompt_as_string = raw_full_prompt.to_string()
    elif isinstance(raw_full_prompt, str): full_prompt_as_string = raw_full_prompt
    else:
        for val in chain_input_dict.values():
            if isinstance(val, str) and "User Question:" in val and "Context Documents:" in val: full_prompt_as_string = val; break
            elif hasattr(val, 'to_string'):
                val_str = val.to_string()
                if "User Question:" in val_str and "Context Documents:" in val_str: full_prompt_as_string = val_str; break
        if not full_prompt_as_string:
            print(f"Warning: 'full_prompt_str' not resolved in llm_pipeline_function from: {chain_input_dict}")
            return "Error: Internal prompt generation problem."
    _p,_,_ = load_persona_data(PERSONA_FILE_PATH)
    msgs = [{"role":"system", "content":f"{_p if _p else PURPOSE} You MUST respond in English. If using <think> block, it must be English."}]
    for turn in chat_history:
        if "query" in turn and "response" in turn:
            msgs.append({"role": "user", "content": str(turn["query"])})
            msgs.append({"role": "assistant", "content": str(turn["response"])})
    msgs.append({"role": "user", "content": full_prompt_as_string})
    payload = {"model":"deepseek-r1-distill-qwen-7b","messages":msgs,"temperature":0.7,"max_tokens":2048,"stream":False}
    try:
        r = requests.post("http://localhost:1234/v1/chat/completions",json=payload,timeout=120)
        r.raise_for_status(); return r.json()["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout: print("Error: LLM API timed out."); return "AI model connection timed out."
    except requests.exceptions.RequestException as e: print(f"LLM API Error: {e}"); return f"AI model connection error: {e}"
    except (KeyError, IndexError, ValueError) as e:
        resp_text = r.text if 'r' in locals() else 'N/A'; print(f"LLM Response Error: {e}. Response: {resp_text[:500]}..."); return "AI model response error."

def create_rag_prompt(bot_name: str, creator_info: str, purpose: str) -> PromptTemplate:
    template_str = f"""You are {bot_name}, a helpful AI assistant for the IIITD college website.
Your persona is friendly, knowledgeable about IIITD based only on the provided context, and strictly focused on assisting with IIITD-related queries.
{creator_info}
{purpose}
**Instructions & Guardrails:**
1.  **Respond in English Only:** All parts of your response, including any internal thought processes or reasoning steps (like those within <think></think> tags if you use them), MUST be in English.
2.  **Prioritize Context:** Base your answers *exclusively* on the provided "Context" documents below.
3.  **Acknowledge Limits:** If the context does not contain the answer, clearly state "Based on the available IIITD documents, I don't have specific information about that."
4.  **Fallback Rule:** Only *after* stating context lacks info, if question is about IIITD but not in context, you *may* use general knowledge *cautiously* and *briefly*, clearly indicating it (e.g., "Generally speaking..."). Do NOT do this for out-of-scope.
5.  **Refuse Out-of-Scope:** Politely decline for requests outside IIITD website documents.
**Context Documents:**
{{context}}
**User Question:**
{{question}}
**Answer ({bot_name}):**"""
    return PromptTemplate(input_variables=["context", "question", "chat_history"], template=template_str)

def format_docs(docs: List[Document]) -> str:
    if not docs: return "No relevant documents found."
    return "\n\n".join([doc.page_content for doc in docs if isinstance(doc, Document)])

def create_contextual_retrieval_query(input_dict: Dict[str, Any], turns_to_include: int) -> str:
    current_query = str(input_dict["question"])
    chat_history: List[Dict[str, str]] = input_dict.get("chat_history", [])
    if not chat_history or turns_to_include <= 0: return current_query
    contextual_parts = []
    for turn_data in chat_history[-turns_to_include:]:
        user_q = str(turn_data.get('query', ''))
        raw_ai_response = str(turn_data.get('response', ''))
        cleaned_ai_response = re.sub(r"<think>.*?</think>\s*|^\*\*Answer:\*\*\s*", "", raw_ai_response, flags=re.DOTALL | re.IGNORECASE).strip()
        if user_q: contextual_parts.append(f"User: {user_q}")
        if cleaned_ai_response: contextual_parts.append(f"Assistant: {cleaned_ai_response}")
    if not contextual_parts: return current_query
    history_str = "\n".join(contextual_parts)
    return f"{history_str}\nUser: {current_query}"

def get_rag_chain(
    bm25_retriever: BM25Retriever,
    reranker: CrossEncoderReranker,
    rag_prompt_template: PromptTemplate,
    history_turns_for_retrieval: int
):
    contextual_query_generator = RunnableLambda(
        lambda x: create_contextual_retrieval_query(x, history_turns_for_retrieval),
        name="ContextualQueryGenerator"
    )

    # This function will be part of the chain to handle BM25 + Reranker
    def retrieve_and_rerank_lambda_func(input_payload: Dict[str, Any]) -> List[Document]:
        # Expects input_payload to be a dictionary containing the 'contextual_query'
        contextual_query_str = input_payload["contextual_query"]
        
        bm25_candidate_docs = bm25_retriever.get_relevant_documents(contextual_query_str)
        if not bm25_candidate_docs:
            return []
        reranked_docs = reranker.compress_documents(
            documents=bm25_candidate_docs,
            query=contextual_query_str
        )
        return reranked_docs

    # The main input to the chain is {"question": str, "chat_history": List[Dict]}
    # Step 1: Generate contextual query.
    # Step 2: Pass this contextual query to the retrieval_rerank_pipeline.
    # Step 3: Combine results with original question and history for the prompt.

    retrieval_rerank_pipeline = (
        RunnableLambda(lambda x: {"contextual_query": contextual_query_generator.invoke(x)}, name="PrepareContextualQueryForRetrieval")
        | RunnableLambda(retrieve_and_rerank_lambda_func, name="RetrieveAndRerank")
    )
    
    prepare_prompt_and_llm_input = RunnableParallel(
        context=(retrieval_rerank_pipeline | RunnableLambda(format_docs, name="FormatDocs")),
        question=RunnableLambda(lambda x: x["question"], name="OriginalQuestion"),
        chat_history=RunnableLambda(lambda x: x.get("chat_history", []), name="ChatHistoryPassthrough")
    )
    
    prepare_for_llm = {
        "full_prompt_str": rag_prompt_template,
        "chat_history": RunnableLambda(lambda x: x["chat_history"], name="ChatHistoryForLLM")
    }

    chain = (
        prepare_prompt_and_llm_input
        | prepare_for_llm
        | RunnableLambda(llm_pipeline_function, name="LLMFunctionCall")
        | StrOutputParser()
    )
    return chain

def initialize_knowledge_base(
    data_root_path: str,
    # device_type: str, # Removed
    corpus_cs: int, corpus_co: int,
    bm25_k: int, cross_encoder_model: str, rerank_n: int,
    persona_file_path_kb: str
) -> Tuple[BM25Retriever | None, CrossEncoderReranker | None, str, dict, dict, PromptTemplate | None]: # Added None types
    print("\n--- Initializing Knowledge Base (BM25 -> CrossEncoder Retriever) ---")
    all_docs = load_documents(data_root_path)
    if not all_docs:
        print("No documents loaded. Aborting."); return None, None, "", {}, {}, None

    try:
        bm25_retriever_obj, reranker_obj = initialize_hybrid_retriever_components(
            all_loaded_docs=all_docs,
            corpus_chunk_size=corpus_cs,
            corpus_chunk_overlap=corpus_co,
            bm25_k_candidates=bm25_k,
            cross_encoder_model_name=cross_encoder_model,
            rerank_top_n=rerank_n
        )
    except Exception as e:
        print(f"Error during hybrid retriever component initialization: {e}")
        return None, None, "", {}, {}, None
        
    if bm25_retriever_obj is None or reranker_obj is None:
        print("Retriever/Reranker initialization failed. Aborting."); return None, None, "", {}, {}, None

    persona_text, faqs_data, paths_data = load_persona_data(persona_file_path_kb)
    rag_prompt_template_obj = create_rag_prompt(BOT_NAME, CREATOR_INFO, PURPOSE)
    print("--- Knowledge Base Initialized ---")
    return bm25_retriever_obj, reranker_obj, persona_text, faqs_data, paths_data, rag_prompt_template_obj

if __name__ == "__main__":
    bm25_retriever, reranker, persona, faqs, paths, rag_prompt_template = initialize_knowledge_base(
        DATA_ROOT,
        # DEVICE, # Removed
        CORPUS_CHUNK_SIZE, CORPUS_CHUNK_OVERLAP,
        BM25_K_CANDIDATES, CROSS_ENCODER_MODEL_NAME, RERANK_TOP_N,
        PERSONA_FILE_PATH
    )

    if bm25_retriever is None or reranker is None or rag_prompt_template is None:
        print("Failed to initialize knowledge base. Exiting.")
    else:
        standard_rag_chain = get_rag_chain(
            bm25_retriever, reranker, rag_prompt_template, HISTORY_TURNS_FOR_RETRIEVAL
        )
        print("\n--- RAG Chain with Hybrid Retrieval & History Test ---")
        session_history_for_test: List[Dict[str, str]] = []
        test_conversation = [
            "What are the admission requirements for the B.Tech CSE program?",
            "What is washkaro app and how does it work?",
            "Who created it? Tell me about Tavpritesh Sethi.",
            "What is the purpose of the app?",
            "How can I access the app?"
        ]
        for i, query in enumerate(test_conversation):
            print(f"\n--- Test Turn {i + 1} ---")
            print(f"❓ User Query: {query}")
            _display_contextual_query = create_contextual_retrieval_query(
                {"question": query, "chat_history": session_history_for_test}, HISTORY_TURNS_FOR_RETRIEVAL)
            print(f"🗣️ Contextual Query for BM25/Reranker (debug): \"{_display_contextual_query}\"")
            print("-" * 20)
            try:
                print(f"💬 Generating response from {BOT_NAME}...")
                chain_input = {"question": query, "chat_history": session_history_for_test}
                final_answer = standard_rag_chain.invoke(chain_input)
                print(f"\n🤖 Answer ({BOT_NAME}):")
                final_answer_cleaned = re.sub(r"<think>.*?</think>\s*|^\*\*Answer:\*\*\s*", "", final_answer, flags=re.DOTALL|re.IGNORECASE).strip()
                print(final_answer_cleaned)
                session_history_for_test.append({"query": query, "response": final_answer})
            except Exception as e:
                print(f"\n💥 ERROR processing query '{query}': {e}")
                import traceback
                traceback.print_exc()
            print("=" * 50)
        print("\n--- All tests completed ---")
        print("Exiting RAG pipeline test.")